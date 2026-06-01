from http.server import BaseHTTPRequestHandler
import cgi
import io
import re
import json
import base64

from docx import Document

_STRIP_PREFIX = re.compile(
    r"^"
    r"(?:"
        r"(?:\[?\d+\]?[\.\):\-]\s*)"
    r")+"
    r"(?=\S)"
)

_TIMESTAMP = re.compile(
    r"^"
    r"(?:"
        r"\d{1,2}:\d{2}(?::\d{2})?"
        r"|"
        r"\d+\s+hours?,\s*\d+\s+minutes?,\s*\d+\s+seconds?"
        r"|"
        r"\d+\s+hours?,\s*\d+\s+minutes?"
        r"|"
        r"\d+\s+hours?,\s*\d+\s+seconds?"
        r"|"
        r"\d+\s+minutes?,\s*\d+\s+seconds?"
        r"|"
        r"\d+\s+hours?"
        r"|"
        r"\d+\s+minutes?"
        r"|"
        r"\d+\s+seconds?"
    r")"
)

_SKIP_LINE              = re.compile(_TIMESTAMP.pattern + r"\s*$")
_STRIP_TIMESTAMP_PREFIX = re.compile(_TIMESTAMP.pattern + r"\s+")
_STRIP_TRAILING         = re.compile(r"\s*sync\s+to\s+video\s+time\s*$", re.IGNORECASE)
_STRIP_BRACKETS         = re.compile(r"\[.*?\]")

def strip_prefix(text: str) -> str:
    text = _STRIP_BRACKETS.sub("", text)
    text = _STRIP_TIMESTAMP_PREFIX.sub("", text)
    text = _STRIP_PREFIX.sub("", text)
    text = _STRIP_TRAILING.sub("", text)
    return text.strip()


def process_paragraphs(raw_lines: list) -> list:
    paragraphs = []
    join_next = False
    for raw in raw_lines:
        if _SKIP_LINE.match(raw.strip()):
            continue
        text = strip_prefix(raw).strip()
        if not text:
            if raw.strip():
                join_next = True
            continue
        first_char = text[0]
        if paragraphs and (join_next or first_char.islower() or first_char.isdigit()):
            paragraphs[-1] = paragraphs[-1].rstrip() + " " + text
        else:
            if paragraphs:
                paragraphs.append("")
            paragraphs.append(text)
        join_next = False
    return paragraphs


def process_docx_bytes(file_bytes: bytes) -> tuple:
    """
    Returns (cleaned_paragraphs, docx_bytes).
    Empty strings in cleaned_paragraphs represent blank lines between paragraphs.
    """
    doc = Document(io.BytesIO(file_bytes))
    raw_lines = [p.text for p in doc.paragraphs if p.text.strip()]
    cleaned = process_paragraphs(raw_lines)

    new_doc = Document()
    try:
        style = doc.styles["Normal"]
        new_style = new_doc.styles["Normal"]
        new_style.font.name = style.font.name
        new_style.font.size = style.font.size
    except Exception:
        pass

    for para_text in cleaned:
        new_doc.add_paragraph(para_text)

    out = io.BytesIO()
    new_doc.save(out)
    out.seek(0)
    return cleaned, out.read()


class handler(BaseHTTPRequestHandler):

    def do_OPTIONS(self):
        self.send_response(200)
        self._cors()
        self.end_headers()

    def do_POST(self):
        try:
            environ = {
                "REQUEST_METHOD": "POST",
                "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
            }
            form = cgi.FieldStorage(fp=self.rfile, headers=self.headers, environ=environ)

            if "file" not in form:
                self._error(400, "No file uploaded.")
                return

            item = form["file"]
            raw_bytes = item.file.read()

            if not raw_bytes:
                self._error(400, "Uploaded file is empty.")
                return

            original_name = getattr(item, "filename", "document.docx") or "document.docx"
            stem = original_name.rsplit(".", 1)[0]

            paragraphs, docx_bytes = process_docx_bytes(raw_bytes)

            body = json.dumps({
                "filename": f"{stem}_filtered.docx",
                "paragraphs": paragraphs,
                "docx_b64": base64.b64encode(docx_bytes).decode()
            }).encode()

            self.send_response(200)
            self._cors()
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        except Exception as exc:
            self._error(500, f"Processing error: {exc}")

    def _cors(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def _error(self, code: int, msg: str):
        body = json.dumps({"error": msg}).encode()
        self.send_response(code)
        self._cors()
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

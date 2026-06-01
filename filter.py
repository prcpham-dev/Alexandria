"""
filter.py – Doc Number Filter
==============================
Reads every .docx file from the "input" folder, strips line numbers and
timestamps, merges continuation lines (lower-case first letter) with the
previous paragraph, and saves the cleaned document to the "output" folder.

Rules applied:
  1. Strip leading numbers / timestamps at the beginning of each line.
     Patterns removed:
       • "1.", "12.", "123." …           (numbered list)
       • "1)", "12)" …                   (numbered list variant)
       • "[1]", "[12]" …                 (bracketed numbers)
       • "00:00", "00:00:00", "1:23:45"  (timestamps)
       • Any combination of the above at the start of a line
  2. After stripping, if the remaining text starts with a LOWER-CASE letter
     → join it to the previous paragraph (continuation).
  3. If it starts with an UPPER-CASE letter (or is a fresh first paragraph)
     → treat it as a new paragraph.
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import copy

# ── Folder paths (relative to this script) ───────────────────────────────────
BASE_DIR   = Path(__file__).parent
INPUT_DIR  = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output"

# ── Regex: strip leading numbers / timestamps ─────────────────────────────────
# Matches things like: "1. ", "2) ", "[3] ", "00:12 ", "1:23:45 ", and combos
_STRIP_PREFIX = re.compile(
    r"^"
    r"(?:"
        r"(?:\[?\d+\]?[\.\):\-]?\s*)"   # number: 1. / 1) / [1] / 1:
        r"|"
        r"(?:\d{1,2}:\d{2}(?::\d{2})?\s*)"  # timestamp: 00:00 / 00:00:00
    r")+"
)


def strip_prefix(text: str) -> str:
    """Remove leading number / timestamp prefixes from a line of text."""
    return _STRIP_PREFIX.sub("", text).strip()


def process_paragraphs(raw_lines: list[str]) -> list[str]:
    """
    Apply the merge/split logic:
      - lower-case first char  → append to previous paragraph
      - upper-case first char  → new paragraph
    Returns a list of clean paragraph strings.
    """
    paragraphs: list[str] = []

    for raw in raw_lines:
        text = strip_prefix(raw).strip()

        # Skip completely empty lines after stripping
        if not text:
            continue

        first_char = text[0]

        if paragraphs and first_char.islower():
            # Continuation – join with a space
            paragraphs[-1] = paragraphs[-1].rstrip() + " " + text
        else:
            # New paragraph (upper-case, digit, punctuation, or very first)
            paragraphs.append(text)

    return paragraphs


def process_docx(src_path: Path, dst_path: Path) -> None:
    """Read src_path .docx, process, write to dst_path."""
    doc = Document(src_path)

    # Collect all non-empty paragraph texts
    raw_lines = [p.text for p in doc.paragraphs if p.text.strip()]

    cleaned = process_paragraphs(raw_lines)

    # Build a new document
    new_doc = Document()

    # Copy core styles / default font from original if possible
    try:
        style = doc.styles["Normal"]
        new_style = new_doc.styles["Normal"]
        new_style.font.name = style.font.name
        new_style.font.size = style.font.size
    except Exception:
        pass

    for para_text in cleaned:
        new_doc.add_paragraph(para_text)

    new_doc.save(dst_path)
    print(f"  ✓  {src_path.name}  →  {dst_path.name}")


def main():
    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)

    docx_files = list(INPUT_DIR.glob("*.docx"))

    if not docx_files:
        print("No .docx files found in the 'input' folder.")
        print(f"  → Put your Word documents in: {INPUT_DIR}")
        return

    print(f"Found {len(docx_files)} file(s) to process...\n")

    for src in docx_files:
        dst = OUTPUT_DIR / src.name
        try:
            process_docx(src, dst)
        except Exception as exc:
            print(f"  ✗  {src.name}  ERROR: {exc}")

    print(f"\nDone! Cleaned files are in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
